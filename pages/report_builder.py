# pages/report_builder.py
"""
Builds the AnalyticaOS .docx report using python-docx.
Pure logic — no Streamlit imports — so it can be unit-tested independently
and reused outside the page layer if needed.

Section functions are added incrementally (Phase 11 Steps 3a-3d).
build_docx_report() is the single public entry point called by pages/report.py.
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pages.executive import LocalExecutiveEngine

# ── Scientific Blue palette, as RGBColor (hex without '#') ───────────────
PRIMARY_RGB   = RGBColor(0x2A, 0x5C, 0x8C)
TEXT_RGB      = RGBColor(0x1B, 0x3A, 0x5C)
MUTED_RGB     = RGBColor(0x5C, 0x72, 0x90)
BORDER_HEX    = "DCE4EC"
SUCCESS_RGB   = RGBColor(0x2E, 0x7D, 0x5B)
WARNING_RGB   = RGBColor(0x8A, 0x63, 0x00)
ERROR_RGB     = RGBColor(0xC0, 0x39, 0x2B)


def _set_cell_shading(cell, hex_color: str):
    """Applies background shading to a table cell (no '#' prefix on hex_color)."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _setup_document() -> Document:
    """Creates a new Document with US Letter page size, 1-inch margins,
    and Arial as the default font for Normal + Heading styles."""
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_RGB

    for level, size, color in [(1, 22, PRIMARY_RGB), (2, 16, PRIMARY_RGB), (3, 13, TEXT_RGB)]:
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Arial'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    return doc


def _add_kpi_table(doc: Document, kpis: dict):
    """Renders a 4-column KPI strip: Health Index, Critical Risks, Tactical Risks, Actions."""
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Corporate Health", "Critical Risks", "Tactical Risks", "Prescriptive Actions"]
    values = [
        f"{kpis['health_score']}/100",
        str(kpis['high_risks']),
        str(kpis['med_risks']),
        str(kpis['action_count']),
    ]

    col_width = Inches(1.6)
    for i in range(4):
        table.columns[i].width = col_width

    for i, (h, v) in enumerate(zip(headers, values)):
        label_cell = table.cell(0, i)
        label_cell.text = h
        label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_cell.paragraphs[0].runs[0]
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = MUTED_RGB
        _set_cell_shading(label_cell, "EFF4FA")

        val_cell = table.cell(1, i)
        val_cell.text = v
        val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_run = val_cell.paragraphs[0].runs[0]
        val_run.font.size = Pt(18)
        val_run.font.bold = True
        val_run.font.color.rgb = PRIMARY_RGB
        _set_cell_shading(val_cell, "FBFCFE")

    doc.add_paragraph()  # spacing after table


def _add_title_page(doc: Document, dataset_name: str, n_rows: int, n_cols: int):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AnalyticaOS")
    run.font.name = 'Arial'
    run.font.size = Pt(34)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_RGB

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Autonomous Data Intelligence Report")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = MUTED_RGB

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if n_rows > 0 and n_cols > 0:
        dims_line = f"Dataset: {dataset_name}  |  {n_rows:,} rows × {n_cols} columns"
    elif n_rows > 0:
        dims_line = f"Dataset: {dataset_name}  |  {n_rows:,} rows"
    else:
        dims_line = f"Dataset: {dataset_name}"

    line1_run = meta.add_run(dims_line)
    line1_run.font.size = Pt(10)
    line1_run.font.color.rgb = MUTED_RGB
    meta.add_run().add_break()  # actual docx line break, not a literal "\n"

    line2_run = meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    line2_run.font.size = Pt(10)
    line2_run.font.color.rgb = MUTED_RGB

    doc.add_page_break()


def _add_executive_summary(doc: Document, briefing: dict):
    doc.add_heading("Executive Summary", level=1)

    narrative = briefing["narrative_summary"].replace("**", "")  # strip markdown bold markers
    p = doc.add_paragraph(narrative)
    p.paragraph_format.space_after = Pt(16)

    _add_kpi_table(doc, briefing["kpis"])


def _add_section_table(doc: Document, headers: list, rows: list, col_widths=None):
    """
    Generic shaded-header table helper shared by Data Quality + EDA sections.
    rows: list of tuples/lists, same length as headers.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr_cells[i], "2A5C8C")  # PRIMARY_RGB as hex
        if col_widths:
            table.columns[i].width = col_widths[i]

    for r_idx, row in enumerate(rows):
        row_cells = table.add_row().cells
        shade = "F5F8FC" if r_idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            row_cells[i].paragraphs[0].runs[0].font.color.rgb = TEXT_RGB
            _set_cell_shading(row_cells[i], shade)

    doc.add_paragraph()  # spacing after table
    return table


def _add_data_quality_section(doc: Document, agent_report: dict):
    """
    Renders the Data Quality section from
    agent_report["results"]["DataQualityAgent"].artifacts:
      missing_summary (df), duplicate_count (int),
      suspicious_type_cols (list), quality_score (0-100)
    """
    doc.add_heading("Data Quality Assessment", level=1)

    dq_result = agent_report.get("results", {}).get("DataQualityAgent")
    if dq_result is None:
        p = doc.add_paragraph(
            "Data Quality Agent did not run for this dataset — "
            "this section has been omitted."
        )
        p.runs[0].font.color.rgb = MUTED_RGB
        p.runs[0].font.italic = True
        return

    artifacts = dq_result.artifacts or {}
    quality_score = artifacts.get("quality_score", 0)
    duplicate_count = artifacts.get("duplicate_count", 0)
    suspicious_cols = artifacts.get("suspicious_type_cols", []) or []
    missing_summary = artifacts.get("missing_summary")

    score_color = SUCCESS_RGB if quality_score >= 80 else (
        WARNING_RGB if quality_score >= 50 else ERROR_RGB
    )

    score_p = doc.add_paragraph()
    score_run = score_p.add_run(f"Data Quality Score: {quality_score}/100")
    score_run.font.bold = True
    score_run.font.size = Pt(13)
    score_run.font.color.rgb = score_color
    score_p.paragraph_format.space_after = Pt(8)

    narrative_parts = []
    narrative_parts.append(
        f"The dataset contains {duplicate_count} duplicate row"
        f"{'s' if duplicate_count != 1 else ''}."
        if duplicate_count else
        "No duplicate rows were detected in the dataset."
    )
    if suspicious_cols:
        narrative_parts.append(
            f"{len(suspicious_cols)} column(s) show inconsistent or "
            f"suspicious data types and may need cleaning: "
            f"{', '.join(suspicious_cols)}."
        )
    else:
        narrative_parts.append("No columns showed suspicious or inconsistent data types.")

    doc.add_paragraph(" ".join(narrative_parts)).paragraph_format.space_after = Pt(10)

    if dq_result.summary:
        summary_p = doc.add_paragraph(dq_result.summary)
        summary_p.paragraph_format.space_after = Pt(10)

    if missing_summary is not None and not missing_summary.empty:
        doc.add_heading("Missing Values by Column", level=2)
        cols = list(missing_summary.columns)
        rows = [tuple(row) for row in missing_summary.itertuples(index=False)]
        _add_section_table(doc, headers=cols, rows=rows)
    else:
        no_missing_p = doc.add_paragraph("No missing values were found in any column.")
        no_missing_p.runs[0].font.color.rgb = SUCCESS_RGB

    findings = dq_result.findings or []
    if findings:
        doc.add_heading("Key Findings", level=2)
        for f in findings:
            severity = f.get("severity", "low")
            sev_color = {"high": ERROR_RGB, "medium": WARNING_RGB, "low": MUTED_RGB}.get(
                severity, MUTED_RGB
            )
            bullet = doc.add_paragraph(style="List Bullet")
            title_run = bullet.add_run(f"{f.get('title', 'Finding')}: ")
            title_run.font.bold = True
            title_run.font.color.rgb = sev_color
            detail_run = bullet.add_run(f.get("detail", ""))
            detail_run.font.color.rgb = TEXT_RGB

    doc.add_paragraph()  # section spacing


def _add_eda_highlights_section(doc: Document, agent_report: dict):
    """
    Renders the EDA Highlights section from
    agent_report["results"]["InsightAgent"].artifacts:
      descriptive_stats (df), strong_correlations (list of tuples),
      moderate_correlations (list of tuples), skewed_cols (list),
      non_normal_cols (list), high_cardinality_cols (list)
    """
    doc.add_heading("Exploratory Data Analysis Highlights", level=1)

    insight_result = agent_report.get("results", {}).get("InsightAgent")
    if insight_result is None:
        p = doc.add_paragraph(
            "Insight Agent did not run for this dataset — "
            "this section has been omitted."
        )
        p.runs[0].font.color.rgb = MUTED_RGB
        p.runs[0].font.italic = True
        return

    artifacts = insight_result.artifacts or {}
    descriptive_stats = artifacts.get("descriptive_stats")
    strong_corr = artifacts.get("strong_correlations", []) or []
    skewed_cols = artifacts.get("skewed_cols", []) or []
    non_normal_cols = artifacts.get("non_normal_cols", []) or []
    high_card_cols = artifacts.get("high_cardinality_cols", []) or []

    if insight_result.summary:
        summary_p = doc.add_paragraph(insight_result.summary)
        summary_p.paragraph_format.space_after = Pt(10)

    if descriptive_stats is not None and not descriptive_stats.empty:
        doc.add_heading("Descriptive Statistics", level=2)
        stats_df = descriptive_stats.reset_index().round(3)
        cols = list(stats_df.columns)
        rows = [tuple(row) for row in stats_df.itertuples(index=False)]
        _add_section_table(doc, headers=cols, rows=rows)

    if strong_corr:
        doc.add_heading("Strong Correlations", level=2)
        rows = []
        for pair in strong_corr:
            # pair expected as (col_a, col_b, correlation_value)
            if len(pair) == 3:
                col_a, col_b, corr_val = pair
                rows.append((col_a, col_b, f"{corr_val:.3f}"))
            else:
                rows.append(tuple(pair))
        _add_section_table(doc, headers=["Variable A", "Variable B", "Correlation"], rows=rows)
    else:
        no_corr_p = doc.add_paragraph("No strong correlations were detected between numeric variables.")
        no_corr_p.runs[0].font.color.rgb = MUTED_RGB

    bullet_groups = [
        ("Skewed Columns", skewed_cols,
         "These columns show notable skew and may benefit from a log or "
         "square-root transformation before modeling."),
        ("Non-Normally Distributed Columns", non_normal_cols,
         "These columns failed normality testing — consider non-parametric "
         "methods when analyzing them."),
        ("High-Cardinality Categorical Columns", high_card_cols,
         "These columns have a large number of unique categories and may "
         "need grouping or encoding strategies before use in models."),
    ]

    for heading, col_list, note in bullet_groups:
        if col_list:
            doc.add_heading(heading, level=2)
            note_p = doc.add_paragraph(note)
            note_p.runs[0].font.italic = True
            note_p.runs[0].font.color.rgb = MUTED_RGB
            for col in col_list:
                doc.add_paragraph(str(col), style="List Bullet")

    findings = insight_result.findings or []
    if findings:
        doc.add_heading("Key Findings", level=2)
        for f in findings:
            severity = f.get("severity", "low")
            sev_color = {"high": ERROR_RGB, "medium": WARNING_RGB, "low": MUTED_RGB}.get(
                severity, MUTED_RGB
            )
            bullet = doc.add_paragraph(style="List Bullet")
            title_run = bullet.add_run(f"{f.get('title', 'Finding')}: ")
            title_run.font.bold = True
            title_run.font.color.rgb = sev_color
            detail_run = bullet.add_run(f.get("detail", ""))
            detail_run.font.color.rgb = TEXT_RGB

    doc.add_paragraph()  # section spacing


def _add_statistical_findings_section(doc: Document, stats_results: list):
    """
    Renders the Statistical Findings section from
    session_state["stats_results"] — a list of dicts written by
    pages/stats.py's _save_stat_result(), schema:
      test_type (str), variables (str), statistic_name (str),
      statistic_value (float), p_value (float), significant (bool),
      interpretation (str)
    """
    doc.add_heading("Statistical Findings", level=1)

    if not stats_results:
        p = doc.add_paragraph(
            "No hypothesis tests were run for this dataset — "
            "this section has been omitted. Visit the Statistical Engine "
            "page to run t-tests, chi-square, ANOVA, or normality tests "
            "before generating the report."
        )
        p.runs[0].font.color.rgb = MUTED_RGB
        p.runs[0].font.italic = True
        doc.add_paragraph()
        return

    n_sig = sum(1 for r in stats_results if r.get("significant"))
    n_total = len(stats_results)

    intro_p = doc.add_paragraph()
    test_verb = "was" if n_total == 1 else "were"
    intro_run = intro_p.add_run(
        f"{n_total} statistical test{'s' if n_total != 1 else ''} "
        f"{test_verb} run on this dataset. {n_sig} of {n_total} returned a "
        f"statistically significant result (p < 0.05)."
    )
    intro_run.font.size = Pt(11)
    intro_p.paragraph_format.space_after = Pt(10)

    headers = ["Test", "Variables", "Statistic", "P-Value", "Verdict"]
    rows = []
    for r in stats_results:
        stat_label = f"{r.get('statistic_name', 'Statistic')} = {r.get('statistic_value', '—')}"
        p_val = r.get("p_value", None)
        p_display = f"{p_val:.4f}" if p_val is not None else "—"
        verdict = "Significant" if r.get("significant") else "Not significant"
        rows.append((
            r.get("test_type", "—"),
            r.get("variables", "—"),
            stat_label,
            p_display,
            verdict,
        ))
    _add_section_table(doc, headers=headers, rows=rows)

    doc.add_heading("Interpretations", level=2)
    for r in stats_results:
        bullet = doc.add_paragraph(style="List Bullet")
        sig_color = SUCCESS_RGB if r.get("significant") else MUTED_RGB
        label_run = bullet.add_run(f"{r.get('test_type', 'Test')} — ")
        label_run.font.bold = True
        label_run.font.color.rgb = sig_color
        detail_run = bullet.add_run(r.get("interpretation", ""))
        detail_run.font.color.rgb = TEXT_RGB

    doc.add_paragraph()  # section spacing


def _add_ml_summary_section(doc: Document, agent_report: dict):
    """
    Renders the ML Summary section from
    agent_report["results"]["ModelingAgent"].artifacts:
      problem_type (str: "regression"/"classification"), target_col (str),
      feature_cols (list), metrics (dict: r2/mae/rmse OR accuracy),
      model (fitted estimator, not rendered), scaler, encoders,
      feature_importance (df)
    """
    doc.add_heading("Machine Learning Summary", level=1)

    ml_result = agent_report.get("results", {}).get("ModelingAgent")
    if ml_result is None:
        p = doc.add_paragraph(
            "Modeling Agent did not run for this dataset — "
            "this section has been omitted."
        )
        p.runs[0].font.color.rgb = MUTED_RGB
        p.runs[0].font.italic = True
        doc.add_paragraph()
        return

    artifacts = ml_result.artifacts or {}
    problem_type = artifacts.get("problem_type", "unknown")
    target_col = artifacts.get("target_col", "—")
    feature_cols = artifacts.get("feature_cols", []) or []
    metrics = artifacts.get("metrics", {}) or {}
    feature_importance = artifacts.get("feature_importance")

    overview_p = doc.add_paragraph()
    overview_run = overview_p.add_run(
        f"A {problem_type} model was trained to predict '{target_col}' "
        f"using {len(feature_cols)} feature{'s' if len(feature_cols) != 1 else ''}."
    )
    overview_run.font.size = Pt(11)
    overview_p.paragraph_format.space_after = Pt(10)

    if ml_result.summary:
        summary_p = doc.add_paragraph(ml_result.summary)
        summary_p.paragraph_format.space_after = Pt(10)

    if metrics:
        doc.add_heading("Model Performance", level=2)
        if problem_type == "regression":
            metric_order = [("r2", "R²"), ("mae", "MAE"), ("rmse", "RMSE")]
        else:
            metric_order = [("accuracy", "Accuracy")]
        # Include any metrics present even if not in the expected order list
        known_keys = {k for k, _ in metric_order}
        extra = [(k, k.upper()) for k in metrics.keys() if k not in known_keys]
        metric_order = metric_order + extra

        headers = [label for _, label in metric_order if metric_order]
        values = []
        for key, label in metric_order:
            val = metrics.get(key)
            values.append(f"{val:.4f}" if isinstance(val, float) else str(val) if val is not None else "—")
        if headers and values:
            _add_section_table(doc, headers=headers, rows=[tuple(values)])

    if feature_importance is not None and not feature_importance.empty:
        doc.add_heading("Feature Importance", level=2)
        top_features = feature_importance.head(10).round(4)
        cols = list(top_features.columns)
        rows = [tuple(row) for row in top_features.itertuples(index=False)]
        _add_section_table(doc, headers=cols, rows=rows)

    findings = ml_result.findings or []
    if findings:
        doc.add_heading("Key Findings", level=2)
        for f in findings:
            severity = f.get("severity", "low")
            sev_color = {"high": ERROR_RGB, "medium": WARNING_RGB, "low": MUTED_RGB}.get(
                severity, MUTED_RGB
            )
            bullet = doc.add_paragraph(style="List Bullet")
            title_run = bullet.add_run(f"{f.get('title', 'Finding')}: ")
            title_run.font.bold = True
            title_run.font.color.rgb = sev_color
            detail_run = bullet.add_run(f.get("detail", ""))
            detail_run.font.color.rgb = TEXT_RGB

    doc.add_paragraph()  # section spacing


def _add_action_matrix_section(doc: Document, briefing: dict):
    """
    Renders the Action Matrix section from briefing["action_matrix"]
    (LocalExecutiveEngine.generate_briefing() output) — a list of dicts:
      {id, action, priority (str: "high"/"medium"/"low"), timeframe, domain}
    """
    doc.add_heading("Prescriptive Action Matrix", level=1)

    action_matrix = briefing.get("action_matrix", []) or []

    if not action_matrix:
        p = doc.add_paragraph(
            "No prescriptive actions were generated for this dataset."
        )
        p.runs[0].font.color.rgb = MUTED_RGB
        p.runs[0].font.italic = True
        doc.add_paragraph()
        return

    priority_order = {"high": 0, "medium": 1, "low": 2}
    sorted_actions = sorted(
        action_matrix,
        key=lambda a: priority_order.get(str(a.get("priority", "low")).lower(), 3),
    )

    intro_p = doc.add_paragraph()
    intro_run = intro_p.add_run(
        f"{len(sorted_actions)} action{'s' if len(sorted_actions) != 1 else ''} "
        f"have been prescribed, ordered by priority."
    )
    intro_run.font.size = Pt(11)
    intro_p.paragraph_format.space_after = Pt(10)

    headers = ["Priority", "Action", "Domain", "Timeframe"]
    rows = [
        (
            str(a.get("priority", "—")).title(),
            a.get("action", "—"),
            a.get("domain", "—"),
            a.get("timeframe", "—"),
        )
        for a in sorted_actions
    ]
    table = _add_section_table(doc, headers=headers, rows=rows)

    # Color the Priority cells by severity for quick visual scanning,
    # overriding the default alternating-row shading on that column only.
    priority_color_map = {"high": ERROR_RGB, "medium": WARNING_RGB, "low": SUCCESS_RGB}
    for r_idx, a in enumerate(sorted_actions, start=1):  # row 0 is the header
        priority_key = str(a.get("priority", "low")).lower()
        color = priority_color_map.get(priority_key, MUTED_RGB)
        cell = table.cell(r_idx, 0)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = color

    doc.add_paragraph()  # section spacing


def _add_appendix_section(doc: Document, aq_roadmap: list, aq_profile: dict, agent_report: dict):
    """
    Renders the Appendix — catch-all for the analysis roadmap and run
    metadata not covered by the named sections above.
    aq_roadmap: list of {step, phase, action, detail, module, priority}
                from Phase 8 Question Generator.
    """
    doc.add_heading("Appendix", level=1)

    doc.add_heading("Analysis Roadmap", level=2)
    if aq_roadmap:
        headers = ["Step", "Phase", "Action", "Module", "Priority"]
        rows = [
            (
                r.get("step", "—"),
                r.get("phase", "—"),
                r.get("action", "—"),
                r.get("module", "—"),
                str(r.get("priority", "—")).title(),
            )
            for r in aq_roadmap
        ]
        _add_section_table(doc, headers=headers, rows=rows)
    else:
        p = doc.add_paragraph(
            "No analysis roadmap was generated for this dataset."
        )
        p.runs[0].font.color.rgb = MUTED_RGB
        p.runs[0].font.italic = True

    doc.add_heading("Dataset Profile", level=2)
    if aq_profile:
        profile_lines = []
        n_rows = aq_profile.get("n_rows")
        n_cols = aq_profile.get("n_cols")
        if n_rows is not None:
            profile_lines.append(f"Rows: {n_rows:,}")
        if n_cols is not None:
            profile_lines.append(f"Columns: {n_cols}")
        for key, label in [
            ("numeric_cols", "Numeric columns"),
            ("cat_cols", "Categorical columns"),
            ("date_cols", "Date columns"),
            ("cols_with_missing", "Columns with missing values"),
        ]:
            val = aq_profile.get(key)
            if val:
                profile_lines.append(f"{label}: {len(val)}")

        if profile_lines:
            for line in profile_lines:
                doc.add_paragraph(line, style="List Bullet")
        else:
            p = doc.add_paragraph("Dataset profile was empty.")
            p.runs[0].font.color.rgb = MUTED_RGB
            p.runs[0].font.italic = True
    else:
        p = doc.add_paragraph("No dataset profile was available for this run.")
        p.runs[0].font.color.rgb = MUTED_RGB
        p.runs[0].font.italic = True

    doc.add_heading("Agent Pipeline Run", level=2)
    pipeline = (agent_report or {}).get("pipeline", [])
    overall_status = (agent_report or {}).get("overall_status", "—")
    if pipeline:
        doc.add_paragraph(f"Pipeline executed: {' → '.join(pipeline)}")
        doc.add_paragraph(f"Overall status: {str(overall_status).title()}")
    else:
        p = doc.add_paragraph("No agent pipeline was executed for this run.")
        p.runs[0].font.color.rgb = MUTED_RGB
        p.runs[0].font.italic = True

    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        f"Report generated by AnalyticaOS on "
        f"{datetime.now().strftime('%B %d, %Y at %H:%M')}."
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = MUTED_RGB
    footer_p.paragraph_format.space_before = Pt(16)


def build_docx_report(context: dict) -> bytes:
    """
    Public entry point. Assembles the full report from session_state context
    (agent_report, aq_profile, aq_roadmap, stats_results) and returns the
    .docx file as in-memory bytes, ready for st.download_button.

    All seven report sections (Executive Summary, Data Quality, EDA
    Highlights, Statistical Findings, ML Summary, Action Matrix, Appendix)
    are implemented as of Phase 11 Step 3d.
    """
    agent_report = context["agent_report"]
    aq_profile   = context["aq_profile"]
    aq_roadmap   = context["aq_roadmap"]
    stats_results = context.get("stats_results", [])

    briefing = LocalExecutiveEngine.generate_briefing(agent_report, aq_profile, aq_roadmap)

    n_rows = aq_profile.get("n_rows", 0) if aq_profile else 0
    n_cols = aq_profile.get("n_cols", 0) if aq_profile else 0

    # Fallback: if aq_profile wasn't populated (Question Generator phase
    # skipped before generating the report) but agents did run, derive a
    # reliable row count from InsightAgent's descriptive_stats (its 'count'
    # column reflects actual non-null row count per column). Column count
    # is NOT guessed here — guessing categorical column counts from
    # artifacts is unreliable and would just replace one wrong number with
    # another. If aq_profile is missing, n_cols stays 0 and the title page
    # falls back to omitting the column count entirely (see _add_title_page).
    if n_rows == 0 and agent_report:
        insight_artifacts = getattr(
            agent_report.get("results", {}).get("InsightAgent"), "artifacts", {}
        ) or {}
        descriptive_stats = insight_artifacts.get("descriptive_stats")
        if descriptive_stats is not None and not descriptive_stats.empty and "count" in descriptive_stats.columns:
            n_rows = int(descriptive_stats["count"].max())

    doc = _setup_document()
    _add_title_page(doc, context.get("dataset_name", "Uploaded Dataset"), n_rows, n_cols)
    _add_executive_summary(doc, briefing)
    _add_data_quality_section(doc, agent_report)
    _add_eda_highlights_section(doc, agent_report)
    _add_statistical_findings_section(doc, stats_results)
    _add_ml_summary_section(doc, agent_report)
    _add_action_matrix_section(doc, briefing)
    _add_appendix_section(doc, aq_roadmap, aq_profile, agent_report)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()